import streamlit as st
import sqlite3
import pandas as pd
from datetime import datetime
from openai import OpenAI
from docx import Document
from io import BytesIO
import hashlib
from pypdf import PdfReader

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="Abogado IA de Bolsillo", 
    page_icon="⚖️",
    layout="centered", # 'centered' se ve mejor en móviles que 'wide'
    initial_sidebar_state="collapsed" # Menú cerrado por defecto en móvil
)

# --- ESTILOS CSS PARA MÓVIL ---
st.markdown("""
<style>
    /* Ajuste para que el chat no se superponga */
    .stChatInput {position: fixed; bottom: 0; padding-bottom: 20px; z-index: 99;}
    /* Botones más grandes para dedos en móvil */
    .stButton button {width: 100%; border-radius: 10px; height: 50px;}
</style>
""", unsafe_allow_html=True)

# --- API KEY ---
try:
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except:
    st.error("⚠️ Falta configurar la API Key en secrets.toml")

# --- BASE DE DATOS ---
def init_db():
    conn = sqlite3.connect('legal_app.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users (username TEXT PRIMARY KEY, password TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS history 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT, date TEXT, filename TEXT, analysis TEXT)''')
    conn.commit()
    conn.close()

def make_hashes(password):
    return hashlib.sha256(str.encode(password)).hexdigest()

def check_hashes(password, hashed_text):
    return make_hashes(password) == hashed_text

def add_user(username, password):
    conn = sqlite3.connect('legal_app.db')
    c = conn.cursor()
    c.execute('INSERT INTO users(username, password) VALUES (?,?)', (username, make_hashes(password)))
    conn.commit()
    conn.close()

def login_user(username, password):
    conn = sqlite3.connect('legal_app.db')
    c = conn.cursor()
    c.execute('SELECT * FROM users WHERE username =? AND password = ?', (username, make_hashes(password)))
    data = c.fetchall()
    conn.close()
    return data

def save_analysis(username, filename, analysis):
    conn = sqlite3.connect('legal_app.db')
    c = conn.cursor()
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    c.execute('INSERT INTO history(username, date, filename, analysis) VALUES (?,?,?,?)', 
              (username, timestamp, filename, analysis))
    conn.commit()
    conn.close()

def get_user_history(username):
    conn = sqlite3.connect('legal_app.db')
    df = pd.read_sql_query("SELECT date, filename, analysis FROM history WHERE username = ? ORDER BY date DESC", conn, params=(username,))
    conn.close()
    return df

# --- LÓGICA DE IA ---

def analyze_general_contract(text_content):
    """Analiza cualquier tipo de documento legal"""
    system_prompt = """
    Actúas como un consultor legal senior experto en legislación argentina.
    Tu tarea es proteger al usuario.
    
    PASO 1: Identifica qué tipo de documento es (Alquiler, Compraventa, Laboral, NDA, Pagaré, etc.).
    PASO 2: Realiza un análisis de riesgos.
    
    Tu respuesta debe tener este formato:
    ### 📂 Tipo de Documento: [Nombre del tipo]
    
    ### 🚦 Nivel de Riesgo: [Bajo/Medio/Alto]
    
    ### ⚠️ Puntos Críticos y Cláusulas Abusivas:
    * [Punto 1]: Explica por qué es riesgoso y cita la ley vigente (CCyC, LCT, etc).
    * [Punto 2]...
    
    ### ✅ Recomendación Final:
    [Consejo directo: Firmar, Negociar o Rechazar]
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o", 
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Analiza este documento: {text_content[:20000]}"}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error en análisis: {e}"

def ask_chat_question(question, contract_text, chat_history):
    """Responde preguntas sobre el contrato cargado"""
    messages = [
        {"role": "system", "content": "Eres un asistente legal. Responde preguntas del usuario basándote EXCLUSIVAMENTE en el contrato que se te proporciona abajo. Si la respuesta no está en el contrato, dilo. Sé breve y claro."},
        {"role": "system", "content": f"CONTRATO: {contract_text[:20000]}"}
    ]
    # Agregar historial de chat para contexto
    messages.extend(chat_history)
    messages.append({"role": "user", "content": question})

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=messages
    )
    return response.choices[0].message.content

def create_docx(analysis_text, filename):
    doc = Document()
    doc.add_heading(f'Análisis Legal: {filename}', 0)
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph(analysis_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFAZ APP ---
def main():
    init_db()
    
    # Manejo de sesión para chat
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "current_contract_text" not in st.session_state:
        st.session_state.current_contract_text = ""
    if "analysis_done" not in st.session_state:
        st.session_state.analysis_done = False

    # SIDEBAR (Navegación)
    with st.sidebar:
        st.title("🏛️ LegalAI")
        if 'logged_in' not in st.session_state:
            st.session_state['logged_in'] = False
            
        if st.session_state['logged_in']:
            st.write(f"Hola, **{st.session_state['username']}**")
            if st.button("Cerrar Sesión"):
                st.session_state['logged_in'] = False
                st.session_state.messages = [] # Limpiar chat
                st.session_state.analysis_done = False
                st.rerun()
            st.markdown("---")
            nav = st.radio("Ir a:", ["Nuevo Análisis", "Historial Guardado"])
        else:
            nav = "Login"

    # PANTALLA LOGIN
    if not st.session_state['logged_in']:
        st.header("Identifícate")
        tab1, tab2 = st.tabs(["Ingresar", "Registrarse"])
        
        with tab1:
            username = st.text_input("Usuario")
            password = st.text_input("Contraseña", type='password')
            if st.button("Entrar", type="primary"):
                if login_user(username, password):
                    st.session_state['logged_in'] = True
                    st.session_state['username'] = username
                    st.rerun()
                else:
                    st.error("Error en datos")
        
        with tab2:
            new_user = st.text_input("Crear Usuario")
            new_pass = st.text_input("Crear Contraseña", type='password')
            if st.button("Registrar"):
                add_user(new_user, new_pass)
                st.success("¡Listo! Ahora inicia sesión.")
        return # Cortamos ejecución aquí si no hay login

    # PANTALLA PRINCIPAL
    if nav == "Nuevo Análisis":
        st.subheader("🤖 Analista Legal Universal")
        st.caption("Sube contratos de alquiler, laborales, servicios, pagarés, etc.")
        
        uploaded_file = st.file_uploader("Sube tu documento (PDF)", type=['pdf', 'txt'])
        
        # Procesar Archivo
        if uploaded_file:
            # Botón de análisis (Solo aparece si no se ha analizado aún)
            if not st.session_state.analysis_done:
                if st.button("🔍 Analizar Documento", type="primary"):
                    with st.spinner("Leyendo y buscando trampas legales..."):
                        text_content = ""
                        if uploaded_file.type == "application/pdf":
                            try:
                                reader = PdfReader(uploaded_file)
                                for page in reader.pages:
                                    text_content += page.extract_text()
                            except:
                                st.error("Error leyendo PDF.")
                        else:
                            text_content = uploaded_file.getvalue().decode("utf-8")
                        
                        if text_content:
                            # 1. Ejecutar análisis
                            analysis = analyze_general_contract(text_content)
                            
                            # 2. Guardar en estado para el chat y UI
                            st.session_state.current_contract_text = text_content
                            st.session_state.current_analysis = analysis
                            st.session_state.analysis_done = True
                            
                            # 3. Guardar en BD
                            save_analysis(st.session_state['username'], uploaded_file.name, analysis)
                            st.rerun()

            # Mostrar Resultados y Chat (Si ya se analizó)
            if st.session_state.analysis_done:
                st.success("✅ Análisis Completado")
                
                # Sección Resumen (Expander para que no ocupe todo el móvil)
                with st.expander("📄 Ver Informe Legal Completo", expanded=True):
                    st.markdown(st.session_state.current_analysis)
                    
                    # Descargar Word
                    docx = create_docx(st.session_state.current_analysis, uploaded_file.name)
                    st.download_button("📥 Bajar informe Word", docx, file_name="analisis.docx")
                
                st.markdown("---")
                st.subheader("💬 Chat con tu Contrato")
                st.caption("Pregunta cosas como: '¿Cuál es la multa si rescindo?' o '¿Es legal la cláusula 5?'")

                # Mostrar historial de chat
                for message in st.session_state.messages:
                    with st.chat_message(message["role"]):
                        st.markdown(message["content"])

                # Input de chat
                if prompt := st.chat_input("Escribe tu duda aquí..."):
                    # 1. Mostrar mensaje usuario
                    st.session_state.messages.append({"role": "user", "content": prompt})
                    with st.chat_message("user"):
                        st.markdown(prompt)

                    # 2. Generar respuesta IA
                    with st.chat_message("assistant"):
                        with st.spinner("Consultando el contrato..."):
                            response = ask_chat_question(
                                prompt, 
                                st.session_state.current_contract_text,
                                st.session_state.messages[:-1] # Historial previo sin el último prompt
                            )
                            st.markdown(response)
                    
                    # 3. Guardar respuesta
                    st.session_state.messages.append({"role": "assistant", "content": response})

                # Botón para limpiar y empezar otro
                if st.button("🔄 Analizar otro documento"):
                    st.session_state.analysis_done = False
                    st.session_state.messages = []
                    st.session_state.current_contract_text = ""
                    st.rerun()

    # PANTALLA HISTORIAL
    elif nav == "Historial Guardado":
        st.subheader("🗄️ Tus análisis anteriores")
        df = get_user_history(st.session_state['username'])
        if not df.empty:
            st.dataframe(df[['date', 'filename', 'analysis']], hide_index=True)
        else:
            st.info("No tienes historial aún.")

if __name__ == '__main__':
    main()
